from collections.abc import Iterator
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document
from fastapi.testclient import TestClient
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)
from sqlalchemy import create_engine, delete
from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.user import User
from app.security.cv_validation import (
    DOCX_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
)

CV_PARSING_TEST_VALUE = "e" * 16


def create_text_pdf(text: str) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )

    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): font,
                }
            ),
        }
    )

    content = DecodedStreamObject()
    content.set_data((f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET").encode("ascii"))
    page[NameObject("/Contents")] = content

    writer.write(output)
    return output.getvalue()


def create_blank_pdf() -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.write(output)
    return output.getvalue()


def create_text_docx() -> bytes:
    output = BytesIO()
    document = Document()

    document.add_paragraph("Jane Doe")
    document.add_paragraph("Backend Software Engineer")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, FastAPI, PostgreSQL"
    table.cell(1, 0).text = "Location"
    table.cell(1, 1).text = "Palestine"

    document.save(output)
    return output.getvalue()


def csrf_headers(client: TestClient) -> dict[str, str]:
    response = client.get("/api/v1/auth/csrf")
    assert response.status_code == 200

    return {
        "X-CSRF-Token": response.json()["csrf_token"],
    }


def register_and_login(
    client: TestClient,
    email: str,
) -> str:
    registration_response = client.post(
        "/api/v1/auth/register",
        json={
            "email": email,
            "display_name": "Parsing User",
            "password": CV_PARSING_TEST_VALUE,
        },
    )
    assert registration_response.status_code == 201

    login_response = client.post(
        "/api/v1/auth/login",
        json={
            "email": email,
            "password": CV_PARSING_TEST_VALUE,
        },
    )
    assert login_response.status_code == 200

    return registration_response.json()["id"]


@pytest.fixture
def created_emails() -> Iterator[list[str]]:
    emails: list[str] = []
    yield emails

    if not emails:
        return

    engine = create_engine(get_settings().database_url)

    with Session(engine) as database:
        database.execute(delete(User).where(User.email.in_(emails)))
        database.commit()

    engine.dispose()


@pytest.fixture(autouse=True)
def clear_parsing_cookies(
    client: TestClient,
) -> Iterator[None]:
    client.cookies.clear()
    yield
    client.cookies.clear()


def test_parsing_requires_authentication_and_csrf(
    client: TestClient,
    created_emails: list[str],
) -> None:
    unauthenticated_response = client.post("/api/v1/cv/parse")
    assert unauthenticated_response.status_code == 401

    email = f"parse-security-{uuid4()}@example.com"
    created_emails.append(email)
    register_and_login(client, email)

    missing_csrf_response = client.post("/api/v1/cv/parse")
    assert missing_csrf_response.status_code == 403

    missing_cv_response = client.post(
        "/api/v1/cv/parse",
        headers=csrf_headers(client),
    )
    assert missing_cv_response.status_code == 404

    missing_text_response = client.get("/api/v1/cv/text")
    assert missing_text_response.status_code == 404


def test_parse_pdf_and_get_private_text(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"parse-pdf-{uuid4()}@example.com"
    created_emails.append(email)

    user_id = register_and_login(client, email)
    headers = csrf_headers(client)

    expected_text = "Jane Doe Software Engineer Python FastAPI PostgreSQL"
    file_data = create_text_pdf(expected_text)

    upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "searchable-resume.pdf",
                file_data,
                PDF_MEDIA_TYPE,
            ),
        },
    )

    assert upload_response.status_code == 200
    assert upload_response.json()["parse_status"] == "pending"

    unavailable_response = client.get("/api/v1/cv/text")
    assert unavailable_response.status_code == 409

    parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )

    assert parse_response.status_code == 200
    assert parse_response.json() == {
        "user_id": user_id,
        "parse_status": "processed",
        "extracted_text": expected_text,
        "character_count": len(expected_text),
    }
    assert parse_response.headers["cache-control"] == "private, no-store"

    text_response = client.get("/api/v1/cv/text")

    assert text_response.status_code == 200
    assert text_response.json() == parse_response.json()
    assert text_response.headers["cache-control"] == "private, no-store"

    metadata_response = client.get("/api/v1/cv")

    assert metadata_response.status_code == 200
    assert metadata_response.json()["parse_status"] == "processed"
    assert "extracted_text" not in metadata_response.json()


def test_parse_docx_paragraphs_and_tables(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"parse-docx-{uuid4()}@example.com"
    created_emails.append(email)

    register_and_login(client, email)
    headers = csrf_headers(client)

    upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "searchable-resume.docx",
                create_text_docx(),
                DOCX_MEDIA_TYPE,
            ),
        },
    )
    assert upload_response.status_code == 200

    parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )

    assert parse_response.status_code == 200

    extracted_text = parse_response.json()["extracted_text"]

    assert "Jane Doe" in extracted_text
    assert "Backend Software Engineer" in extracted_text
    assert "Skills | Python, FastAPI, PostgreSQL" in extracted_text
    assert "Location | Palestine" in extracted_text


def test_failed_extraction_updates_parse_status(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"parse-empty-{uuid4()}@example.com"
    created_emails.append(email)

    register_and_login(client, email)
    headers = csrf_headers(client)

    upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "scanned-resume.pdf",
                create_blank_pdf(),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert upload_response.status_code == 200

    parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )

    assert parse_response.status_code == 422
    assert parse_response.json() == {"detail": "No searchable text was found in the CV."}

    metadata_response = client.get("/api/v1/cv")

    assert metadata_response.status_code == 200
    assert metadata_response.json()["parse_status"] == "failed"

    text_response = client.get("/api/v1/cv/text")
    assert text_response.status_code == 409


def test_reupload_resets_extracted_text(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"parse-reset-{uuid4()}@example.com"
    created_emails.append(email)

    register_and_login(client, email)
    headers = csrf_headers(client)

    first_text = "First Resume Python Developer"
    second_text = "Second Resume DevOps Engineer"

    first_upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "first-resume.pdf",
                create_text_pdf(first_text),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert first_upload_response.status_code == 200

    first_parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )
    assert first_parse_response.status_code == 200
    assert first_parse_response.json()["extracted_text"] == first_text

    second_upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "second-resume.pdf",
                create_text_pdf(second_text),
                PDF_MEDIA_TYPE,
            ),
        },
    )

    assert second_upload_response.status_code == 200
    assert second_upload_response.json()["parse_status"] == "pending"

    old_text_response = client.get("/api/v1/cv/text")
    assert old_text_response.status_code == 409

    second_parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )

    assert second_parse_response.status_code == 200
    assert second_parse_response.json()["extracted_text"] == second_text


def test_cv_skills_require_authentication_and_processed_cv(
    client: TestClient,
    created_emails: list[str],
) -> None:
    unauthenticated_response = client.get("/api/v1/cv/skills")
    assert unauthenticated_response.status_code == 401

    email = f"skills-security-{uuid4()}@example.com"
    created_emails.append(email)

    register_and_login(client, email)

    missing_cv_response = client.get("/api/v1/cv/skills")
    assert missing_cv_response.status_code == 404
    assert missing_cv_response.json() == {"detail": "CV not found."}

    upload_response = client.post(
        "/api/v1/cv",
        headers=csrf_headers(client),
        files={
            "file": (
                "pending-resume.pdf",
                create_text_pdf("Python FastAPI Developer"),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert upload_response.status_code == 200

    pending_response = client.get("/api/v1/cv/skills")

    assert pending_response.status_code == 409
    assert pending_response.json() == {
        "detail": "CV skills are not available.",
    }


def test_parsing_persists_and_returns_normalized_skills(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"skills-extraction-{uuid4()}@example.com"
    created_emails.append(email)

    user_id = register_and_login(client, email)
    headers = csrf_headers(client)

    upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "skills-resume.pdf",
                create_text_pdf("Python FastAPI postgres React.js JavaScript"),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert upload_response.status_code == 200

    parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )
    assert parse_response.status_code == 200

    skills_response = client.get("/api/v1/cv/skills")

    assert skills_response.status_code == 200
    assert skills_response.json() == {
        "user_id": user_id,
        "skills": [
            "Python",
            "FastAPI",
            "PostgreSQL",
            "React",
            "JavaScript",
        ],
        "skill_count": 5,
        "extraction_version": "taxonomy-v1",
    }
    assert skills_response.headers["cache-control"] == "private, no-store"


def test_cv_skills_allow_empty_results(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"skills-empty-{uuid4()}@example.com"
    created_emails.append(email)

    user_id = register_and_login(client, email)
    headers = csrf_headers(client)

    upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "no-skills-resume.pdf",
                create_text_pdf("Jane Doe Product Support Specialist"),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert upload_response.status_code == 200

    parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )
    assert parse_response.status_code == 200

    skills_response = client.get("/api/v1/cv/skills")

    assert skills_response.status_code == 200
    assert skills_response.json() == {
        "user_id": user_id,
        "skills": [],
        "skill_count": 0,
        "extraction_version": "taxonomy-v1",
    }


def test_reupload_resets_extracted_skills(
    client: TestClient,
    created_emails: list[str],
) -> None:
    email = f"skills-reset-{uuid4()}@example.com"
    created_emails.append(email)

    register_and_login(client, email)
    headers = csrf_headers(client)

    first_upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "first-skills.pdf",
                create_text_pdf("Python FastAPI Developer"),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert first_upload_response.status_code == 200

    first_parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )
    assert first_parse_response.status_code == 200

    first_skills_response = client.get("/api/v1/cv/skills")
    assert first_skills_response.status_code == 200
    assert first_skills_response.json()["skills"] == [
        "Python",
        "FastAPI",
    ]

    second_upload_response = client.post(
        "/api/v1/cv",
        headers=headers,
        files={
            "file": (
                "second-skills.pdf",
                create_text_pdf("Java Spring Boot Docker Engineer"),
                PDF_MEDIA_TYPE,
            ),
        },
    )
    assert second_upload_response.status_code == 200
    assert second_upload_response.json()["parse_status"] == "pending"

    reset_response = client.get("/api/v1/cv/skills")
    assert reset_response.status_code == 409

    second_parse_response = client.post(
        "/api/v1/cv/parse",
        headers=headers,
    )
    assert second_parse_response.status_code == 200

    second_skills_response = client.get("/api/v1/cv/skills")

    assert second_skills_response.status_code == 200
    assert second_skills_response.json()["skills"] == [
        "Java",
        "Spring Boot",
        "Docker",
    ]
    assert second_skills_response.json()["skill_count"] == 3
