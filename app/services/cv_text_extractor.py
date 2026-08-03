from io import BytesIO
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from lxml.etree import XMLSyntaxError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from app.security.cv_validation import (
    DOCX_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
)

MAX_PDF_PAGES = 30
MAX_PDF_CONTENT_STREAM_SIZE = 10 * 1024 * 1024
MAX_EXTRACTED_CHARACTERS = 100_000


class CVTextExtractionError(ValueError):
    """Raised when searchable text cannot be safely extracted."""


def extract_cv_text(
    file_data: bytes,
    media_type: str,
) -> str:
    try:
        if media_type == PDF_MEDIA_TYPE:
            raw_text = _extract_pdf_text(file_data)
        elif media_type == DOCX_MEDIA_TYPE:
            raw_text = _extract_docx_text(file_data)
        else:
            raise CVTextExtractionError("The CV media type is unsupported.")
    except CVTextExtractionError:
        raise
    except (
        BadZipFile,
        KeyError,
        PackageNotFoundError,
        PdfReadError,
        TypeError,
        ValueError,
        XMLSyntaxError,
    ) as error:
        raise CVTextExtractionError("The CV text could not be extracted.") from error

    normalized_text = _normalize_text(raw_text)

    if not normalized_text:
        raise CVTextExtractionError("No searchable text was found in the CV.")

    if len(normalized_text) > MAX_EXTRACTED_CHARACTERS:
        raise CVTextExtractionError("The extracted CV text exceeds the allowed limit.")

    return normalized_text


def _extract_pdf_text(file_data: bytes) -> str:
    reader = PdfReader(
        BytesIO(file_data),
        strict=False,
        root_object_recovery_limit=1000,
    )

    if reader.is_encrypted:
        raise CVTextExtractionError("Encrypted PDF files are not supported.")

    if len(reader.pages) > MAX_PDF_PAGES:
        raise CVTextExtractionError(f"A CV cannot contain more than {MAX_PDF_PAGES} pages.")

    extracted_pages: list[str] = []
    extracted_character_count = 0

    for page in reader.pages:
        contents = page.get_contents()

        if contents is not None:
            content_stream = contents.get_data()

            if len(content_stream) > MAX_PDF_CONTENT_STREAM_SIZE:
                raise CVTextExtractionError("A PDF content stream exceeds the safe limit.")

        page_text = page.extract_text() or ""
        extracted_pages.append(page_text)
        extracted_character_count += len(page_text)

        if extracted_character_count > MAX_EXTRACTED_CHARACTERS:
            raise CVTextExtractionError("The extracted CV text exceeds the allowed limit.")

    return "\n".join(extracted_pages)


def _extract_docx_text(file_data: bytes) -> str:
    document = Document(BytesIO(file_data))
    text_parts: list[str] = []

    for paragraph in document.paragraphs:
        text_parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            text_parts.append(row_text)

    return "\n".join(text_parts)


def _normalize_text(raw_text: str) -> str:
    normalized_lines: list[str] = []

    for line in raw_text.replace("\x00", "").splitlines():
        normalized_line = " ".join(line.replace("\xa0", " ").split())

        if normalized_line:
            normalized_lines.append(normalized_line)

    return "\n".join(normalized_lines)
