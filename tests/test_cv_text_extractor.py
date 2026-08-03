from io import BytesIO

import pytest
from docx import Document
from pypdf import PdfWriter
from pypdf.generic import (
    DecodedStreamObject,
    DictionaryObject,
    NameObject,
)

from app.security.cv_validation import (
    DOCX_MEDIA_TYPE,
    PDF_MEDIA_TYPE,
)
from app.services.cv_text_extractor import (
    MAX_PDF_PAGES,
    CVTextExtractionError,
    extract_cv_text,
)


def create_text_pdf(text: str) -> bytes:
    output = BytesIO()
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)

    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )

    page[NameObject("/Resources")] = DictionaryObject(
        {
            NameObject("/Font"): DictionaryObject(
                {
                    NameObject("/F1"): font,
                }
            ),
        }
    )

    content = DecodedStreamObject()
    content.set_data((f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET").encode("ascii"))
    page[NameObject("/Contents")] = content

    writer.write(output)
    return output.getvalue()


def create_blank_pdf(page_count: int = 1) -> bytes:
    output = BytesIO()
    writer = PdfWriter()

    for _ in range(page_count):
        writer.add_blank_page(width=612, height=792)

    writer.write(output)
    return output.getvalue()


def create_text_docx() -> bytes:
    output = BytesIO()
    document = Document()

    document.add_paragraph("Jane Doe")
    document.add_paragraph("  Software   Engineer  ")

    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python and FastAPI"

    document.save(output)
    return output.getvalue()


def test_extracts_text_from_pdf() -> None:
    file_data = create_text_pdf("Jane Doe Software Engineer Python FastAPI")

    extracted_text = extract_cv_text(
        file_data,
        PDF_MEDIA_TYPE,
    )

    assert extracted_text == ("Jane Doe Software Engineer Python FastAPI")


def test_extracts_paragraphs_and_tables_from_docx() -> None:
    file_data = create_text_docx()

    extracted_text = extract_cv_text(
        file_data,
        DOCX_MEDIA_TYPE,
    )

    assert "Jane Doe" in extracted_text
    assert "Software Engineer" in extracted_text
    assert "Skills | Python and FastAPI" in extracted_text


def test_rejects_pdf_without_searchable_text() -> None:
    file_data = create_blank_pdf()

    with pytest.raises(
        CVTextExtractionError,
        match="No searchable text",
    ):
        extract_cv_text(file_data, PDF_MEDIA_TYPE)


def test_rejects_pdf_with_too_many_pages() -> None:
    file_data = create_blank_pdf(MAX_PDF_PAGES + 1)

    with pytest.raises(
        CVTextExtractionError,
        match="cannot contain more",
    ):
        extract_cv_text(file_data, PDF_MEDIA_TYPE)


def test_rejects_unsupported_media_type() -> None:
    with pytest.raises(
        CVTextExtractionError,
        match="unsupported",
    ):
        extract_cv_text(b"content", "text/plain")
